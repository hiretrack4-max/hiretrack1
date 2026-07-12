"""
Raw text extraction from resume files (BRD Module 2 -> Module 3).

Supported formats (BRD Module 2): PDF, DOCX, and legacy DOC.

* PDF  -> pdfplumber
* DOCX -> python-docx
* DOC  -> best-effort binary scrape (legacy Word binary format has no pure-Python
          reader in the dependency set); if nothing usable is recovered we fail
          gracefully with a clear message instead of crashing.

Every failure mode (missing dependency, encrypted / empty / corrupt file) is
surfaced as a typed :class:`ExtractionError` carrying a human-readable message
that the caller stores in ``Resume.parse_error``.
"""
from __future__ import annotations

import re
from pathlib import Path

# Resume.FileType values, lower-cased, mapped to the extractor to use.
PDF = "PDF"
DOC = "DOC"
DOCX = "DOCX"


class ExtractionError(Exception):
    """Raised when raw text cannot be extracted from a resume file."""


def _extract_pdf(path: str) -> str:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - dependency guaranteed in prod
        raise ExtractionError(
            "pdfplumber is not installed; cannot extract text from PDF resumes."
        ) from exc

    try:
        parts: list[str] = []
        with pdfplumber.open(path) as pdf:
            if not pdf.pages:
                raise ExtractionError("PDF contains no pages.")
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text:
                    parts.append(text)
    except ExtractionError:
        raise
    except Exception as exc:  # encrypted / corrupt / unreadable
        raise ExtractionError(
            f"Could not read PDF (it may be encrypted or corrupt): {exc}"
        ) from exc

    return "\n".join(parts).strip()


def _extract_docx(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "python-docx is not installed; cannot extract text from DOCX resumes."
        ) from exc

    try:
        document = docx.Document(path)
    except Exception as exc:
        raise ExtractionError(
            f"Could not read DOCX (it may be corrupt or password protected): {exc}"
        ) from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Include table cell text (skills matrices, contact tables, etc.).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


# Runs of >= 4 printable characters recovered from the raw byte stream.
_PRINTABLE_RUN = re.compile(rb"[\x20-\x7e]{4,}")


def _extract_doc(path: str) -> str:
    """Best-effort extraction for the legacy binary .doc format.

    There is no pure-Python .doc reader in the dependency set. We first try
    python-docx in case the file is actually a mislabeled .docx, then fall back
    to scraping printable ASCII runs out of the binary. If that yields nothing
    meaningful, we fail gracefully with a clear message.
    """
    # Some uploads are .docx mislabeled as .doc — try the real reader first.
    try:
        return _extract_docx(path)
    except ExtractionError:
        pass

    try:
        raw = Path(path).read_bytes()
    except OSError as exc:
        raise ExtractionError(f"Could not open DOC file: {exc}") from exc

    runs = [m.group().decode("ascii", "ignore") for m in _PRINTABLE_RUN.finditer(raw)]
    text = "\n".join(runs).strip()

    if len(text) < 40:
        raise ExtractionError(
            "Legacy .doc format could not be parsed. Please re-upload the resume "
            "as PDF or DOCX."
        )
    return text


_EXTRACTORS = {
    PDF: _extract_pdf,
    DOCX: _extract_docx,
    DOC: _extract_doc,
}


def extract_text(path: str, file_type: str) -> str:
    """Extract raw text from ``path`` given its ``Resume.FileType`` value.

    Raises :class:`ExtractionError` on any failure, including an empty result.
    """
    key = (file_type or "").upper()
    extractor = _EXTRACTORS.get(key)
    if extractor is None:
        # Fall back on the filename extension when file_type is absent/unknown.
        ext = Path(path).suffix.lstrip(".").upper()
        extractor = _EXTRACTORS.get("DOCX" if ext == "DOCX" else ext)
    if extractor is None:
        raise ExtractionError(f"Unsupported resume format: {file_type or path!r}.")

    text = extractor(path)
    if not text or not text.strip():
        raise ExtractionError(
            "No text could be extracted from the resume (it may be a scanned "
            "image or an empty document)."
        )
    return text
